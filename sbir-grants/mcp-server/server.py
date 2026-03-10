"""
SBIR Data MCP Server
專注於經濟部統計處官方 API

功能：
1. 經濟部統計處總體統計資料庫 API
2. 工研院 IEK、資策會 MIC 由 Claude 的 search_web 處理
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx import Document
from save_extracted_answers import MCP_save_extracted_answers
from proposal_state import MCP_save_generated_section, MCP_export_proposal_to_word, MCP_get_all_saved_sections
from company_verify import MCP_verify_company_eligibility_by_g0v
from quality_check import MCP_check_proposal_quality
from enrich_answer import MCP_enrich_answer
from ingest_reference_document import MCP_ingest_reference_document, MCP_read_document_for_tagging, MCP_ingest_tagged_chunks, MCP_retrieve_reference_chunks
from section_generation_prompt import MCP_get_section_generation_prompt
from ai_draft_review import MCP_get_ai_draft_review_prompt
import os
import glob
import re
import time
import math
import subprocess
import logging
from typing import Any
from datetime import datetime
from pathlib import Path

from mcp.server import Server
from mcp.types import Tool, TextContent
from pydantic import BaseModel

logger = logging.getLogger(__name__)

# Import proposal generator functions
try:
    from proposal_generator_impl import (
        start_proposal_generator,
        save_answer,
        get_progress,
        generate_proposal,
        STATE_FILE
    )
except ImportError as e:
    logger.warning(f"proposal_generator_impl 匯入失敗: {e}")
    # 定義空的 fallback 函式，避免 NameError

    async def start_proposal_generator(*args, **kwargs):  # type: ignore
        return [TextContent(type="text", text="❌ proposal_generator_impl 模組未找到，請確認檔案存在。")]

    async def save_answer(*args, **kwargs):  # type: ignore
        return [TextContent(type="text", text="❌ proposal_generator_impl 模組未找到。")]

    async def get_progress(*args, **kwargs):  # type: ignore
        return [TextContent(type="text", text="❌ proposal_generator_impl 模組未找到。")]

    async def generate_proposal(*args, **kwargs):  # type: ignore
        return [TextContent(type="text", text="❌ proposal_generator_impl 模組未找到。")]
    STATE_FILE = None  # type: ignore


# ============================================
# 資料模型
# ============================================

class MOEAStatData(BaseModel):
    """經濟部統計處數據格式"""
    category: str        # 類別
    period: str          # 統計期間
    value: float         # 數值
    unit: str            # 單位
    source_url: str      # 來源網址

# ============================================
# MCP Server 初始化
# ============================================


app = Server("sbir-data-server")

# ============================================
# 工具定義
# ============================================


@app.list_tools()
async def list_tools() -> list[Tool]:
    """定義可用的工具"""
    return [
        Tool(
            name="save_extracted_answers",
            description="將結構化的萃取內容存入本地 SQLite 以支援 Smart Extraction Pipeline。",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_id": {"type": "string", "description": "專案的 ID。"},
                    "section_id": {"type": "string", "description": "章節的 ID (例如 section_1)。"},
                    "answers": {"type": "array", "items": {"type": "string"}, "description": "由 Claude 提取的回答陣列。"}
                },
                "required": ["project_id", "section_id", "answers"]
            }
        ),
        Tool(
            name="get_section_generation_prompt",
            description="針對特定的章節，獲取專家 Persona 與詳細生成指示 (Section-level Chunk Generation)。",
            inputSchema={
                "type": "object",
                "properties": {
                    "section_id": {"type": "string", "description": "章節的 ID (例如 section_1)。"}
                },
                "required": ["section_id"]
            }
        ),
        Tool(
            name="ingest_reference_document",
            description="將參考文件切塊並加標籤，存入本地的向量或 SQLite 系統中 (Local RAG Tagging Engine)。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "欲輸入的文檔路徑。"},
                    "tags": {"type": "array", "items": {"type": "string"}, "description": "要標註的 SBIR 章節列表。"}
                },
                "required": ["file_path", "tags"]
            }
        ),
        Tool(
            name="read_document_for_tagging",
            description="讀取並切分檔案，讓 Claude 可以先讀過各個段落，並手動判斷/賦予各段落適合的 SBIR 章節標籤 (AI Auto-Tagging 階段一)。會回傳 JSON 陣列，包含各段的 chunk_index 與 content。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "欲讀取並進行智能分塊的文檔路徑（支援 PDF, docx, md, txt 等）。"}
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="ingest_tagged_chunks",
            description="將已打好專屬標籤的段落寫入知識庫 (AI Auto-Tagging 階段二)。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "來源文檔的相對路徑。"},
                    "tagged_chunks": {
                        "type": "string",
                        "description": "JSON 格式字串，例如： '[{\"chunk_index\": 0, \"tags\": [\"section_1\"]}, {\"chunk_index\": 1, \"tags\": [\"section_2\"]}]'"
                    }
                },
                "required": ["file_path", "tagged_chunks"]
            }
        ),
        Tool(
            name="retrieve_reference_chunks",
            description="[撰寫段落必備] 從使用者的私有知識庫中，撈出打有特定標籤（例如 section_1）的所有語意參考段落。讓你能基於這些具體素材撰寫計畫書。",
            inputSchema={
                "type": "object",
                "properties": {
                    "tags": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "要查詢的章節標籤列表，例如 [\"section_1\"]"
                    }
                }
            }
        ),
        Tool(
            name="verify_company_eligibility_by_g0v",
            description="透過 g0v 台灣公司登記 API，自動檢核使用者公司的 SBIR 申請資格（成立狀況、資本額小於一億、無外資等）。",
            inputSchema={
                "type": "object",
                "properties": {
                    "company_name": {"type": "string", "description": "欲檢驗的公司完整名稱"},
                    "capital_from_user": {"type": "string", "description": "使用者問卷填寫的資本額（例如 \"500萬元\"），以便無法解析 g0v 資本額時參考"},
                    "employee_size_from_user": {"type": "string", "description": "使用者問卷填寫的員工人數（例如 \"10人\"）"}
                },
                "required": ["company_name"]
            }
        ),
        Tool(
            name="save_generated_section",
            description="將 Claude 為使用者撰寫好的正式計畫書段落儲存到本機資料庫。寫完一個章節就應該呼叫一次。",
            inputSchema={
                "type": "object",
                "properties": {
                    "section_index": {"type": "integer", "description": "章節編號（例如 1, 2, 3）"},
                    "title": {"type": "string", "description": "章節標題"},
                    "content": {"type": "string", "description": "撰寫完成的 Markdown 內容"}
                },
                "required": ["section_index", "title", "content"]
            }
        ),
        Tool(
            name="export_proposal_to_word",
            description="將所有曾經透過 save_generated_section 儲存的草稿段落，依序組裝並匯出成完整的 Word (.docx) 檔供使用者下載。",
            inputSchema={
                "type": "object",
                "properties": {
                    "output_filename": {"type": "string", "description": "欲匯出的檔名，預設為 SBIR_Proposal_Draft.docx"}
                }
            }
        ),
        Tool(
            name="get_all_saved_sections",
            description="將所有已儲存的計畫書段落合併為一個大型字串並回傳。此工具專門用於當需要生成 Pitch Deck 簡報，或者需要全盤審視整份計畫書脈絡時呼叫。",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        Tool(
            name="search_knowledge_base",
            description="搜尋 SBIR 知識庫中的相關文件。可搜尋方法論、FAQ、檢核清單、案例等。",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜尋關鍵字，如：創新、市場分析、經費、資格等"
                    },
                    "category": {
                        "type": "string",
                        "description": "文件類別（可選）",
                        "enum": ["methodology", "faq", "checklist", "case_study", "template", "all"],
                        "default": "all"
                    }
                },
                "required": ["query"]
            }
        ),
        Tool(
            name="read_document",
            description="讀取 SBIR 知識庫中的特定文件內容",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文件的相對路徑，如：references/methodology_innovation.md"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="query_moea_statistics",
            description="查詢經濟部統計處總體統計資料庫（官方 API）。可查詢產業產值、出口、就業等數據。",
            inputSchema={
                "type": "object",
                "properties": {
                    "industry": {
                        "type": "string",
                        "description": "產業別，如：機械、化工、電子、資通訊"
                    },
                    "stat_type": {
                        "type": "string",
                        "description": "統計類型：產值、出口、就業人數",
                        "enum": ["產值", "出口", "就業人數"]
                    },
                    "start_year": {
                        "type": "integer",
                        "description": "起始年份（西元年）",
                        "default": 2020
                    },
                    "end_year": {
                        "type": "integer",
                        "description": "結束年份（西元年）",
                        "default": 2024
                    }
                },
                "required": ["industry", "stat_type"]
            }
        ),
        Tool(
            name="search_moea_website",
            description="搜尋經濟部統計處網站（當 API 無法滿足需求時使用）",
            inputSchema={
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "搜尋關鍵字"
                    }
                },
                "required": ["keyword"]
            }
        ),
        Tool(
            name="start_proposal_generator",
            description="開始互動式計畫書生成器，載入問題並初始化狀態",
            inputSchema={
                "type": "object",
                "properties": {
                    "phase": {
                        "type": "string",
                        "description": "計畫階段",
                        "enum": ["phase1", "phase2"],
                        "default": "phase1"
                    }
                },
                "required": []
            }
        ),
        Tool(
            name="save_answer",
            description="保存問答答案到狀態檔案",
            inputSchema={
                "type": "object",
                "properties": {
                    "question_id": {
                        "type": "string",
                        "description": "問題 ID"
                    },
                    "answer": {
                        "type": "string",
                        "description": "用戶的答案"
                    }
                },
                "required": ["question_id", "answer"]
            }
        ),
        Tool(
            name="get_progress",
            description="取得計畫書生成進度",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        Tool(
            name="generate_proposal",
            description="根據已回答的問題生成完整計畫書",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        Tool(
            name="update_knowledge_base",
            description="更新 SBIR 知識庫到最新版本（從 GitHub 拉取更新）",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        Tool(
            name="check_proposal",
            description="檢核 SBIR 計畫書完整度。這是自我檢查工具，用來確認計畫書是否涵蓋所有必要內容，非評審結果預測。",
            inputSchema={
                "type": "object",
                "properties": {
                    "proposal_content": {
                        "type": "string",
                        "description": "計畫書內容（全文或主要章節）"
                    },
                    "phase": {
                        "type": "string",
                        "description": "計畫階段",
                        "enum": ["phase1", "phase2"],
                        "default": "phase1"
                    }
                },
                "required": ["proposal_content"]
            }
        ),
        Tool(
            name="calculate_budget",
            description="SBIR 經費試算工具。根據計畫階段和總經費，自動建議各項經費分配比例。",
            inputSchema={
                "type": "object",
                "properties": {
                    "phase": {
                        "type": "string",
                        "description": "計畫階段",
                        "enum": ["phase1", "phase2", "phase2plus"],
                        "default": "phase1"
                    },
                    "total_budget": {
                        "type": "number",
                        "description": "計畫總經費（萬元）"
                    },
                    "project_type": {
                        "type": "string",
                        "description": "計畫類型",
                        "enum": ["技術研發", "軟體開發", "硬體開發", "服務創新"],
                        "default": "技術研發"
                    }
                },
                "required": ["total_budget"]
            }
        ),
        Tool(
            name="export_proposal_word",
            description="將計畫書匯出為 Word 檔案（.docx）。支援完整格式，可直接送件。",
            inputSchema={
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "計畫書的 Markdown 內容"
                    },
                    "filename": {
                        "type": "string",
                        "description": "檔案名稱（不含副檔名）",
                        "default": "SBIR_計畫書"
                    },
                    "company_name": {
                        "type": "string",
                        "description": "公司名稱（用於頁首）"
                    },
                    "project_name": {
                        "type": "string",
                        "description": "計畫名稱（用於標題頁）"
                    }
                },
                "required": ["content"]
            }
        ),
        Tool(
            name="calculate_roi",
            description="計算建議的產值目標。根據補助金額、產業別、計畫階段等，自動計算最低、建議、優秀三種產值目標及 ROAS。",
            inputSchema={
                "type": "object",
                "properties": {
                    "subsidy_amount": {
                        "type": "number",
                        "description": "補助金額（萬元）"
                    },
                    "industry": {
                        "type": "string",
                        "description": "產業別",
                        "enum": ["製造業", "機械", "化工/材料", "電子", "資通訊", "軟體", "數位服務", "生技/醫療", "服務業", "服務創新"],
                        "default": "製造業"
                    },
                    "phase": {
                        "type": "string",
                        "description": "計畫階段",
                        "enum": ["phase1", "phase2"],
                        "default": "phase1"
                    },
                    "company_revenue": {
                        "type": "number",
                        "description": "公司年營收（萬元，可選）",
                        "default": 0
                    }
                },
                "required": ["subsidy_amount"]
            }
        ),
        Tool(
            name="validate_roi",
            description="驗證產值是否合理。檢查預期產值與補助金額的比例（ROAS）是否符合產業標準，並提供改進建議。",
            inputSchema={
                "type": "object",
                "properties": {
                    "subsidy_amount": {
                        "type": "number",
                        "description": "補助金額（萬元）"
                    },
                    "expected_revenue_3years": {
                        "type": "number",
                        "description": "預期 3 年累積產值（萬元）"
                    },
                    "industry": {
                        "type": "string",
                        "description": "產業別",
                        "enum": ["製造業", "機械", "化工/材料", "電子", "資通訊", "軟體", "數位服務", "生技/醫療", "服務業", "服務創新"],
                        "default": "製造業"
                    },
                    "phase": {
                        "type": "string",
                        "description": "計畫階段",
                        "enum": ["phase1", "phase2"],
                        "default": "phase1"
                    }
                },
                "required": ["subsidy_amount", "expected_revenue_3years"]
            }
        ),
        Tool(
            name="enrich_answer",
            description="評估使用者的問卷回答是否達到 SBIR 審查標準，並提供具體的改善建議與擴寫提示。適合在使用者完成每個問題後自動呼叫，確保品質達標。",
            inputSchema={
                "type": "object",
                "properties": {
                    "question_id": {
                        "type": "string",
                        "description": "問題的 ID（如 problem_description, solution_description 等）"
                    },
                    "user_answer": {
                        "type": "string",
                        "description": "使用者的回答內容"
                    },
                    "question_text": {
                        "type": "string",
                        "description": "問題原文（可選，用於顯示）"
                    },
                    "context": {
                        "type": "object",
                        "description": "目前已收集到的其他答案，可選。用於 team_experience、經費與營收等題目的輔助整理。"
                    }
                },
                "required": ["question_id", "user_answer"]
            }
        ),
        Tool(
            name="check_proposal_quality",
            description="對已完成的 SBIR 計畫書草稿進行 6 維度品質審查，包含：創新差異化、市場三層分析、商業模式產值、執行期程、數據可信度、整體專業度。提供完整報告與改善建議。",
            inputSchema={
                "type": "object",
                "properties": {
                    "proposal_text": {
                        "type": "string",
                        "description": "計畫書全文（可留空，系統將嘗試從生成器狀態自動讀取）"
                    }
                },
                "required": []
            }
        ),
        Tool(
            name="get_ai_draft_review_prompt",
            description="產生針對特定草稿章節的 AI 審閱 (Track Changes) 提示詞。會自動帶入此專案的事實基準 (Ground Truth) 與過件關鍵清單，打擊幻覺並強制擴寫。請直接將此 Prompt 作為 System Prompt 或 User Prompt 交由 LLM 執行。",
            inputSchema={
                "type": "object",
                "properties": {
                    "section_index": {
                        "type": "integer",
                        "description": "欲審閱的章節編號 (例如：1 為公司簡介，2 為問題陳述)"
                    }
                },
                "required": ["section_index"]
            }
        )
    ]

# ============================================
# 工具執行
# ============================================


@app.call_tool()
async def call_tool(name: str, arguments: Any) -> list[TextContent]:
    """執行工具"""
    if name == "save_extracted_answers":
        res = await MCP_save_extracted_answers(arguments["project_id"], arguments["section_id"], arguments["answers"])
        return [TextContent(type="text", text=res)]
    elif name == "get_section_generation_prompt":
        res = await MCP_get_section_generation_prompt(arguments["section_id"])
        return [TextContent(type="text", text=res)]
    elif name == "ingest_reference_document":
        res = await MCP_ingest_reference_document(arguments["file_path"], arguments.get("tags", []))
        return [TextContent(type="text", text=res)]
    elif name == "read_document_for_tagging":
        res = await MCP_read_document_for_tagging(arguments["file_path"])
        return [TextContent(type="text", text=res)]
    elif name == "ingest_tagged_chunks":
        res = await MCP_ingest_tagged_chunks(arguments["file_path"], arguments["tagged_chunks"])
        return [TextContent(type="text", text=res)]
    elif name == "retrieve_reference_chunks":
        res = await MCP_retrieve_reference_chunks(arguments.get("tags"))
        return [TextContent(type="text", text=res)]
    elif name == "verify_company_eligibility_by_g0v":
        res = await MCP_verify_company_eligibility_by_g0v(
            arguments["company_name"],
            arguments.get("capital_from_user"),
            arguments.get("employee_size_from_user")
        )
        return [TextContent(type="text", text=res)]
    elif name == "save_generated_section":
        res = await MCP_save_generated_section(arguments["section_index"], arguments["title"], arguments["content"])
        return [TextContent(type="text", text=res)]
    elif name == "export_proposal_to_word":
        res = await MCP_export_proposal_to_word(arguments.get("output_filename", "SBIR_Proposal_Draft.docx"))
        return [TextContent(type="text", text=res)]
    elif name == "get_all_saved_sections":
        res = await MCP_get_all_saved_sections()
        return [TextContent(type="text", text=res)]
    elif name == "enrich_answer":
        res = await MCP_enrich_answer(
            arguments["question_id"],
            arguments["user_answer"],
            arguments.get("question_text", ""),
            arguments.get("context")
        )
        return [TextContent(type="text", text=res)]
    elif name == "check_proposal_quality":
        res = await MCP_check_proposal_quality(arguments.get("proposal_text", ""))
        return [TextContent(type="text", text=res)]
    elif name == "search_knowledge_base":
        res = await search_knowledge_base(
            arguments["query"],
            arguments.get("category", "all")
        )
        return [TextContent(type="text", text=str(res))] if not isinstance(res, list) else res
    elif name == "read_document":
        return await read_document(arguments["file_path"])  # type: ignore
    elif name == "query_moea_statistics":
        return await query_moea_statistics(  # type: ignore
            arguments["industry"],
            arguments["stat_type"],
            arguments.get("start_year", 2020),
            arguments.get("end_year", 2024)
        )
    elif name == "search_moea_website":
        return await search_moea_website(arguments["keyword"])  # type: ignore
    elif name == "start_proposal_generator":
        return await start_proposal_generator(arguments.get("phase", "phase1"))
    elif name == "save_answer":
        return await save_answer(arguments["question_id"], arguments["answer"])
    elif name == "get_progress":
        return await get_progress()
    elif name == "generate_proposal":
        return await generate_proposal()
    elif name == "update_knowledge_base":
        return await update_knowledge_base()
    elif name == "check_proposal":
        return await check_proposal(
            arguments["proposal_content"],
            arguments.get("phase", "phase1")
        )
    elif name == "calculate_budget":
        return await calculate_budget(
            arguments["total_budget"],
            arguments.get("phase", "phase1"),
            arguments.get("project_type", "技術研發")
        )
    elif name == "export_proposal_word":
        return await export_proposal_word(
            arguments["content"],
            arguments.get("filename", "SBIR_計畫書"),
            arguments.get("company_name", ""),
            arguments.get("project_name", "")
        )
    elif name == "calculate_roi":
        return await call_calculate_roi(
            arguments["subsidy_amount"],
            arguments.get("industry", "製造業"),
            arguments.get("phase", "phase1"),
            arguments.get("company_revenue", 0)
        )
    elif name == "validate_roi":
        return await call_validate_roi(
            arguments["subsidy_amount"],
            arguments["expected_revenue_3years"],
            arguments.get("industry", "製造業"),
            arguments.get("phase", "phase1")
        )
    elif name == "get_ai_draft_review_prompt":
        res = await MCP_get_ai_draft_review_prompt(arguments["section_index"])
        return [TextContent(type="text", text=res)]
    else:
        raise ValueError(f"Unknown tool: {name}")

# ============================================
# 核心功能：知識庫搜尋與讀取
# ============================================

# 取得專案根目錄（server.py 的上一層）
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# 版本檢查（每天最多檢查一次）
LAST_VERSION_CHECK = 0.0
VERSION_CHECK_INTERVAL = 86400  # 24 小時


def check_for_updates() -> str | None:
    """
    檢查是否有新版本可用
    返回更新提醒訊息，如果已是最新則返回 None
    """
    global LAST_VERSION_CHECK

    current_time = time.time()

    # 每 24 小時只檢查一次
    if current_time - LAST_VERSION_CHECK < VERSION_CHECK_INTERVAL:
        return None

    LAST_VERSION_CHECK = current_time

    try:
        # 取得本地最新 commit
        local_result = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=PROJECT_ROOT,
            capture_output=True,
            text=True,
            timeout=5
        )
        if local_result.returncode != 0:
            return None
        local_commit = local_result.stdout.strip()[:7]

        # 取得遠端最新 commit
        subprocess.run(
            ["git", "fetch", "--quiet"],
            cwd=PROJECT_ROOT,
            capture_output=True,
            timeout=10
        )

        remote_result = subprocess.run(
            ["git", "rev-parse", "origin/main"],
            cwd=PROJECT_ROOT,
            capture_output=True,
            text=True,
            timeout=5
        )
        if remote_result.returncode != 0:
            return None
        remote_commit = remote_result.stdout.strip()[:7]

        # 比較版本
        if local_commit != remote_commit:
            return f"\n\n---\n💡 **有新版本可用！** 您的版本：`{local_commit}`，最新版本：`{remote_commit}`\n請說「**更新知識庫**」來獲得最新內容。"

        return None

    except Exception:
        # 任何錯誤都靜默忽略
        return None


async def search_knowledge_base(query: str, category: str = "all") -> str:
    """
    搜尋知識庫
    混合搜尋：關鍵字 + RAG 語意搜尋
    """

    # ===== 0. 檢查快取 =====
    from search_cache import get_cache
    cache = get_cache()
    cached_result = cache.get(query, category)
    if cached_result:
        return cached_result + "\n\n💡 *此結果來自快取，回應速度更快*"

    # 定義搜尋目錄
    search_dirs = {
        "methodology": "references/methodology_*.md",
        "faq": "faq/*.md",
        "checklist": "checklists/*.md",
        "case_study": "examples/case_studies/*.md",
        "template": "templates/*.md",
        "all": "**/*.md"
    }

    pattern = search_dirs.get(category, "**/*.md")
    search_path = os.path.join(PROJECT_ROOT, pattern)

    # 搜尋檔案
    files = glob.glob(search_path, recursive=True)

    # ===== 1. 關鍵字搜尋（含同義詞擴展）=====
    from query_expansion import get_expanded_keywords
    keywords = get_expanded_keywords(query)
    keyword_results = {}  # path -> score

    for file_path in files:
        file_name = os.path.basename(file_path).lower()
        relative_path = os.path.relpath(file_path, PROJECT_ROOT)

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().lower()

            score = 0
            matched_keywords = []

            for keyword in keywords:
                if keyword in file_name:
                    score += 3
                    matched_keywords.append(keyword)
                elif keyword in content:
                    count = min(content.count(keyword), 5)
                    score += count
                    matched_keywords.append(keyword)

            if score > 0:
                keyword_results[relative_path] = {
                    "path": relative_path,
                    "name": os.path.basename(file_path),
                    "category": get_category_from_path(relative_path),
                    "keyword_score": score,
                    "matched_keywords": len(set(matched_keywords)),
                    "total_keywords": len(keywords)
                }
        except (OSError, UnicodeDecodeError) as e:
            logger.debug(f"無法讀取檔案 {file_path}: {e}")
            continue

    # ===== 2. 語意搜尋 (RAG) =====
    semantic_results = {}  # path -> {similarity, content, metadata}
    semantic_available = False

    try:
        from vector_search import semantic_search, needs_reindex, rerank_results, mmr_sort

        persist_dir = os.path.join(os.path.dirname(__file__), "chroma_db")

        if not needs_reindex(persist_dir):
            semantic_available = True
            results = semantic_search(query, persist_dir, n_results=15)

            for result in results:
                semantic_results[result["id"]] = {
                    "similarity": result["similarity"],
                    "content": result.get("content", ""),
                    "metadata": result.get("metadata", {})
                }
    except Exception as e:
        # 語意搜尋不可用，僅使用關鍵字搜尋
        logger.warning(f"語意搜尋不可用: {e}")

    # ===== 3. 混合排序 =====
    KEYWORD_WEIGHT = 0.4
    SEMANTIC_WEIGHT = 0.6

    all_paths = set(keyword_results.keys()) | set(semantic_results.keys())

    # 正規化關鍵字分數
    max_keyword_score = max([float(str(r["keyword_score"])) for r in keyword_results.values()], default=1.0)

    final_scores = []
    for path in all_paths:
        # 關鍵字分數（正規化到 0-1）
        kw_info = keyword_results.get(path, {})
        kw_score = float(str(kw_info.get("keyword_score", 0))) / max_keyword_score if max_keyword_score > 0 else 0

        # 語意分數（已經是 0-1）
        sem_info = semantic_results.get(path, {})
        sem_score = float(str(sem_info.get("similarity", 0))) if isinstance(sem_info, dict) else 0.0

        # 加權總分
        if semantic_available:
            final_score = kw_score * KEYWORD_WEIGHT + sem_score * SEMANTIC_WEIGHT
        else:
            final_score = kw_score  # 只有關鍵字

        # 取得文件資訊
        if path in keyword_results:
            info = keyword_results[path].copy()
        else:
            # 從語意結果取得 metadata
            sem_metadata = sem_info.get("metadata", {}) if isinstance(sem_info, dict) else {}
            info = {
                "path": sem_metadata.get("file_path", path),
                "name": sem_metadata.get("file", os.path.basename(path)),
                "category": get_category_from_path(path),
                "matched_keywords": 0,
                "total_keywords": len(keywords)
            }

        info["final_score"] = final_score
        info["semantic_score"] = sem_score

        # 從語意結果取得內容預覽
        if isinstance(sem_info, dict):
            content = sem_info.get("content", "")
            metadata = sem_info.get("metadata", {})

            if metadata.get("preview"):
                info["preview"] = metadata.get("preview")

            if content:
                # Bug Y1 fix: was 100 chars (barely one sentence). Increased to 1000 to give Claude
                # meaningful context from each search result rather than a fragment.
                info["content_snippet"] = content[:1000].replace('\n', ' ').strip()

            # 提取來源資訊
            if metadata.get("source_url"):
                info["source_url"] = metadata.get("source_url")
            if metadata.get("source_title"):
                info["source_title"] = metadata.get("source_title")
            if metadata.get("source_date"):
                info["source_date"] = metadata.get("source_date")

        final_scores.append(info)

    # ===== 3.5. 先進行 Re-ranking (對前 20 名) =====
    # 只有當 semantic_available 為真時才進行，因為需要模型
    if semantic_available and len(final_scores) > 0:
        # 取前 20 名進行重排序
        top_candidates = final_scores[:20]
        remaining = final_scores[20:]

        # 準備 Re-ranking 需要的格式 (需有 content)
        # 注意：keyword search 結果可能沒有 content，需要處理
        for cand in top_candidates:
            if "content" not in cand:
                # 嘗試讀取部分內容
                try:
                    with open(os.path.join(str(PROJECT_ROOT), str(cand["path"])), 'r', encoding='utf-8') as f:
                        cand["content"] = f.read(1000)  # 只讀前 1000 字
                except (OSError, UnicodeDecodeError):
                    cand["content"] = cand["name"]  # 降級使用檔名

        # 執行 Re-ranking
        try:
            reranked = rerank_results(query, top_candidates, top_k=20)
            final_scores = reranked + remaining
        except Exception as e:
            print(f"Re-ranking 步驟錯誤: {e}")

    # ===== 3.6. 時間加權 =====

    def apply_time_weight(score: float, source_date: str) -> float:
        """對較新的文件給予更高權重"""
        if not source_date:
            return score

        try:
            # 解析日期
            if isinstance(source_date, str):
                date_obj = datetime.strptime(source_date, "%Y-%m-%d")
            else:
                date_obj = source_date

            # 計算天數差異
            days_old = (datetime.now() - date_obj).days

            # 時間衰減因子（1年衰減到 37%）
            time_decay = math.exp(-days_old / 365)

            # 加權：70% 原始分數 + 30% 時間因素
            # 注意：re-rank 分數可能是負的 logit，這裡的加權公式可能需要調整
            # 簡單起見，如果是 rerank_score，直接加分
            weighted_score = score * (0.7 + 0.3 * time_decay)

            return weighted_score
        except (ValueError, TypeError):
            return score

    # 應用時間加權
    for info in final_scores:
        if source_date := info.get("source_date"):
            # 優先使用 rerank_score，如果沒有則使用 final_score
            target_score_key = "rerank_score" if "rerank_score" in info else "final_score"
            info[target_score_key] = apply_time_weight(float(str(info[target_score_key])), str(source_date))

    # ===== 3.7. MMR 多樣性排序 =====
    if semantic_available and len(final_scores) > 0:
        try:
            final_scores = mmr_sort(final_scores, lambda_param=0.7)
        except Exception as e:
            print(f"MMR 步驟錯誤: {e}")
            # 降級：按分數排序
            final_scores.sort(key=lambda x: float(str(x.get("rerank_score", x.get("final_score", 0)))), reverse=True)
    else:
        # 僅按分數排序
        final_scores.sort(key=lambda x: float(str(x.get("rerank_score", x.get("final_score", 0)))), reverse=True)

    # ===== 4. 格式化結果 =====
    if not final_scores:
        result = f"""
## 搜尋結果

找不到與「{query}」相關的文件。

**建議**：
- 試試其他關鍵字
- 查看完整文件列表：README.md
"""
    else:
        search_mode = "🔍 混合搜尋（關鍵字 + AI 語意）" if semantic_available else "🔍 關鍵字搜尋"
        result = f"""
## 搜尋結果：找到 {len(final_scores)} 個相關段落

**搜尋模式**：{search_mode}
**搜尋關鍵字**：{query}

💡 **提示**：以下結果包含文件來源和內容預覽，Claude 會自動閱讀這些內容並為您綜合答案。

"""
        for i, file_info in enumerate(final_scores[:10], 1):
            # 顯示匹配度
            if semantic_available and float(str(file_info.get("semantic_score", 0))) > 0:
                relevance = f"相關度: {float(str(file_info['final_score']))*100:.0f}%"
            else:
                match_ratio = f"{file_info.get('matched_keywords', 0)}/{file_info['total_keywords']}"
                relevance = f"匹配: {match_ratio} 關鍵字"

            # 檢查是否有 chunk 預覽
            preview = file_info.get("preview", "")
            content_snippet = file_info.get("content_snippet", "")
            source_url = file_info.get("source_url")
            source_title = file_info.get("source_title")
            source_date = file_info.get("source_date")

            result += f"{i}. **{file_info['name']}** ({relevance})\n"

            if preview:
                result += f"   > 📄 {preview}\n"

            if content_snippet:
                result += f"   > 「{content_snippet}」\n"

            result += f"   - 📁 類別：{file_info['category']}\n"
            result += f"   - 📍 位置：`{file_info['path']}`\n"

            # 顯示官方來源
            if source_url:
                result += f"   - 🔗 **官方出處**：{source_url}\n"
            if source_title:
                result += f"   - 📋 來源標題：{source_title}\n"
            if source_date:
                result += f"   - 📅 發布日期：{source_date}\n"

            result += "   - 🔍 使用 `read_document` 工具可讀取完整內容\n\n"

        if len(final_scores) > 10:
            result += f"\n（還有 {len(final_scores) - 10} 個相關段落未顯示）\n"

        # ===== 5. 搜尋建議 =====
        try:
            from search_suggestions import generate_suggestions
            suggestions = generate_suggestions(query, final_scores)

            if suggestions:
                result += "\n💡 **您可能也想了解**：\n"
                for sugg in suggestions:
                    # 這裡使用特殊的 markdown 連結格式讓 Claude 知道這是建議查詢
                    # 格式：[查詢: 問題](command:search_knowledge_base?query=問題)
                    # 但 Claude 不一定支援 command link，直接列出文字即可
                    result += f"- {sugg}\n"
        except Exception as e:
            print(f"搜尋建議生成失敗: {e}")

        if not semantic_available:
            result += "\n💡 **提示**：執行 `python mcp-server/build_index.py` 可啟用 AI 語意搜尋，提升搜尋準確度。\n"

        # 加入引用說明
        result += "\n---\n\n"
        result += "📌 **如何使用這些結果**：\n"
        result += "- Claude 會自動閱讀上述內容並為您綜合答案\n"
        result += "- 答案會包含具體的來源引用\n"
        result += "- 如需查證，可使用 `read_document` 工具閱讀完整文件\n"

    # 寫回快取
    cache.set(query, category, result)

    # 檢查是否有新版本
    update_notice = check_for_updates()
    if update_notice:
        result += update_notice

    return result


async def read_document(file_path: str) -> list[TextContent]:
    """
    讀取指定的文件內容
    """

    full_path = os.path.join(PROJECT_ROOT, file_path)

    # 安全檢查：確保路徑在專案目錄內（使用 realpath 防止符號連結穿越）
    if not os.path.realpath(full_path).startswith(os.path.realpath(PROJECT_ROOT)):
        return [TextContent(
            type="text",
            text="❌ 錯誤：無法讀取專案目錄外的檔案"
        )]

    # 檢查檔案是否存在
    if not os.path.exists(full_path):
        return [TextContent(
            type="text",
            text=f"❌ 錯誤：找不到檔案 `{file_path}`\n\n請使用 `search_knowledge_base` 工具搜尋正確的檔案路徑。"
        )]

    # 讀取檔案
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            content = f.read()

        result = f"""
## 📄 {os.path.basename(file_path)}

**路徑**：`{file_path}`

---

{content}
"""
        return [TextContent(type="text", text=result)]

    except Exception as e:
        return [TextContent(
            type="text",
            text=f"❌ 讀取檔案失敗：{str(e)}"
        )]


def get_category_from_path(path: str) -> str:
    """根據路徑判斷文件類別"""
    if "methodology" in path:
        return "方法論"
    elif "faq" in path:
        return "常見問題"
    elif "checklist" in path:
        return "檢核清單"
    elif "case_studies" in path:
        return "案例研究"
    elif "template" in path:
        return "範本"
    elif "quick_start" in path:
        return "快速啟動"
    else:
        return "其他"

# ============================================
# 核心功能：查詢經濟部統計處 API
# ============================================


async def query_moea_statistics(
    industry: str,
    stat_type: str,
    start_year: int,
    end_year: int
) -> list[TextContent]:
    """
    查詢經濟部統計處總體統計資料庫 API

    API 文件：https://nstatdb.dgbas.gov.tw/dgbasAll/webMain.aspx?sys=100&funid=API
    """

    # 產業代碼對應表（需要根據實際 API 文件調整）
    industry_codes = {
        "機械": "C29",
        "化工": "C20",
        "電子": "C26",
        "資通訊": "C26",
        "生技": "C21",
        "服務業": "G-S"
    }

    # stat_type_codes 映射表（不需要已定義，直接在查詢中使用 stat_type）
    _ = {
        "產値": "production",
        "出口": "export",
        "就業人數": "employment"
    }

    industry_code = industry_codes.get(industry)
    if not industry_code:
        return [TextContent(
            type="text",
            text=f"❌ 不支援的產業別：{industry}\n\n支援的產業：{', '.join(industry_codes.keys())}"
        )]

    # 注意：經濟部統計處 API 需要「功能代碼」才能查詢，目前返回指引訊息及書籤替代方案
    result = f"""
## 經濟部統計處查詢結果

**產業別**：{industry}（代碼：{industry_code}）
**統計類型**：{stat_type}
**查詢期間**：{start_year} - {end_year}

---

⚠️ **API 實作說明**：

經濟部統計處提供總體統計資料庫 API，但需要：
1. 查詢「功能代碼」（每個統計表有唯一代碼）
2. 功能代碼列表：https://nstatdb.dgbas.gov.tw/

**建議替代方案**：
由於功能代碼查詢複雜，建議使用 Claude 的 `search_web` 工具：

```
search_web("{industry} {stat_type} site:dgbas.gov.tw OR site:moea.gov.tw")
```

**API 查詢範例**（需要功能代碼）：
```
https://nstatdb.dgbas.gov.tw/dgbasAll/webMain.aspx?sys=100&funid=API
  ?function=[功能代碼]
  &startTime={start_year}-01
  &endTime={end_year}-12
```

---

**來源**：
- 經濟部統計處：https://www.moea.gov.tw/Mns/dos/
- 總體統計資料庫：https://nstatdb.dgbas.gov.tw/
"""

    return [TextContent(type="text", text=result)]

# ============================================
# 輔助功能：搜尋經濟部網站
# ============================================


async def search_moea_website(keyword: str) -> list[TextContent]:
    """提供搜尋建議（實際搜尋由 Claude 的 search_web 執行）"""

    result = f"""
## 經濟部統計處搜尋建議

**搜尋關鍵字**：{keyword}

---

**建議使用 Claude 的 `search_web` 工具**：

```
search_web("{keyword} site:dgbas.gov.tw OR site:moea.gov.tw")
```

**推薦查詢網站**：
- 經濟部統計處：https://www.moea.gov.tw/Mns/dos/
- 總體統計資料庫：https://nstatdb.dgbas.gov.tw/
- 產業統計：https://www.moea.gov.tw/Mns/dos/content/SubMenu.aspx?menu_id=6730

**查詢技巧**：
- 加上年份：`{keyword} 2024`
- 指定統計類型：`{keyword} 產值` 或 `{keyword} 出口`
"""

    return [TextContent(type="text", text=result)]

# ============================================
# Server 啟動
# ============================================
# 知識庫更新功能
# ============================================


async def update_knowledge_base() -> list[TextContent]:
    """
    從 GitHub 拉取最新版本的知識庫
    """
    try:
        # 執行 git pull
        result = subprocess.run(
            ["git", "pull"],
            cwd=PROJECT_ROOT,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode == 0:
            output = result.stdout.strip()
            if "Already up to date" in output or "已經是最新" in output:
                return [TextContent(
                    type="text",
                    text="✅ **知識庫已是最新版本！**\n\n您的 SBIR Skill 知識庫已經是最新的了，無需更新。"
                )]
            else:
                return [TextContent(
                    type="text",
                    text=f"✅ **知識庫更新成功！**\n\n已從 GitHub 拉取最新版本。\n\n更新內容：\n```\n{output}\n```\n\n請重新啟動 Claude Desktop 以載入新內容。"
                )]
        else:
            error_msg = result.stderr.strip() or result.stdout.strip()
            return [TextContent(
                type="text",
                text=f"❌ **更新失敗**\n\n錯誤訊息：\n```\n{error_msg}\n```\n\n可能的原因：\n1. 沒有網路連線\n2. 專案目錄不是用 git clone 下載的\n3. 有未提交的本地修改\n\n您可以手動執行：\n```bash\ncd {PROJECT_ROOT} && git pull\n```"
            )]

    except subprocess.TimeoutExpired:
        return [TextContent(
            type="text",
            text="❌ **更新超時**\n\n網路連線可能太慢，請稍後再試或手動執行：\n```bash\ngit pull\n```"
        )]
    except FileNotFoundError:
        return [TextContent(
            type="text",
            text="❌ **找不到 Git**\n\n您的系統可能沒有安裝 Git，或 Git 不在系統路徑中。\n\n請手動下載最新版本：\nhttps://github.com/backtrue/sbir-grants/archive/refs/heads/main.zip"
        )]
    except Exception as e:
        return [TextContent(
            type="text",
            text=f"❌ **更新失敗**\n\n發生未預期的錯誤：{str(e)}\n\n請手動執行：\n```bash\ncd {PROJECT_ROOT} && git pull\n```"
        )]

# ============================================
# 經費試算功能
# ============================================


async def calculate_budget(total_budget: float, phase: str = "phase1", project_type: str = "技術研發") -> list[TextContent]:
    """
    根據計畫階段和類型，建議經費分配比例
    """

    # 驗證經費範圍
    phase_limits = {
        "phase1": {"max": 150, "subsidy_max": 150, "name": "Phase 1"},
        "phase2": {"max": 2400, "subsidy_max": 1200, "name": "Phase 2"},
        "phase2plus": {"max": 1200, "subsidy_max": 600, "name": "Phase 2+"}
    }

    limit = phase_limits.get(phase, phase_limits["phase1"])

    if float(total_budget) > float(str(limit["max"])):
        return [TextContent(
            type="text",
            text=f"⚠️ **經費超過上限**\n\n{limit['name']} 計畫總經費上限為 {limit['max']} 萬元，您輸入的是 {total_budget} 萬元\n\n（補助上限：{limit['subsidy_max']} 萬元）"
        )]

    # 根據計畫類型調整比例
    allocation_templates = {
        "技術研發": {
            "人事費": {"ratio": 0.40, "desc": "研發人員薪資"},
            "消耗性器材": {"ratio": 0.20, "desc": "材料、試劑、零組件"},
            "設備費": {"ratio": 0.15, "desc": "研發設備採購或租用"},
            "委託研究費": {"ratio": 0.10, "desc": "委外測試、認證"},
            "差旅費": {"ratio": 0.05, "desc": "技術交流、客戶訪談"},
            "專利費": {"ratio": 0.03, "desc": "專利申請與維護"},
            "管理費": {"ratio": 0.07, "desc": "行政管理費用"}
        },
        "軟體開發": {
            "人事費": {"ratio": 0.55, "desc": "工程師薪資"},
            "消耗性器材": {"ratio": 0.05, "desc": "開發工具"},
            "雲端服務費": {"ratio": 0.15, "desc": "雲端主機、API 費用"},
            "委託研究費": {"ratio": 0.10, "desc": "委外設計、測試"},
            "差旅費": {"ratio": 0.05, "desc": "客戶訪談、技術交流"},
            "專利費": {"ratio": 0.03, "desc": "軟體著作權"},
            "管理費": {"ratio": 0.07, "desc": "行政管理費用"}
        },
        "硬體開發": {
            "人事費": {"ratio": 0.35, "desc": "研發人員薪資"},
            "消耗性器材": {"ratio": 0.25, "desc": "電子零件、材料"},
            "設備費": {"ratio": 0.20, "desc": "量測設備、打樣"},
            "委託研究費": {"ratio": 0.08, "desc": "委外測試、認證"},
            "差旅費": {"ratio": 0.04, "desc": "供應商拜訪"},
            "專利費": {"ratio": 0.03, "desc": "專利申請"},
            "管理費": {"ratio": 0.05, "desc": "行政管理費用"}
        },
        "服務創新": {
            "人事費": {"ratio": 0.50, "desc": "服務開發人員"},
            "消耗性器材": {"ratio": 0.08, "desc": "服務所需材料"},
            "場地費": {"ratio": 0.12, "desc": "服務場域租用"},
            "委託研究費": {"ratio": 0.12, "desc": "市場調查、顧問"},
            "差旅費": {"ratio": 0.08, "desc": "客戶訪談"},
            "行銷費": {"ratio": 0.05, "desc": "推廣活動"},
            "管理費": {"ratio": 0.05, "desc": "行政管理費用"}
        }
    }

    template = allocation_templates.get(project_type, allocation_templates["技術研發"])

    # 計算補助金額
    subsidy = min(float(total_budget) * 0.5, float(limit["subsidy_max"]))  # type: ignore
    self_fund = total_budget - subsidy

    # 生成經費分配表
    output = f"""# 💰 SBIR 經費試算結果

## 基本資訊

| 項目 | 金額（萬元） |
|------|-------------|
| 計畫總經費 | **{total_budget:,.0f}** |
| 補助款（50%） | **{subsidy:,.0f}** |
| 自籌款（50%） | **{self_fund:,.0f}** |

> 計畫階段：{limit['name']}
> 計畫類型：{project_type}

---

## 建議經費分配

| 項目 | 比例 | 金額（萬元） | 說明 |
|------|------|-------------|------|
"""

    for item_name, item_data in template.items():
        amount = float(total_budget) * float(str(item_data["ratio"]))  # type: ignore
        output += f"| {item_name} | {int(float(str(item_data['ratio']))*100)}% | {amount:,.0f} | {item_data['desc']} |\n"

    output += """
---

## ⚠️ 注意事項

1. **人事費上限**：原則上不超過總經費 50%
2. **管理費上限**：不超過總經費 10%
3. **設備費限制**：Phase 1 盡量避免大型設備採購

## 📋 經費編列建議

"""

    # 根據計畫類型給予建議
    if project_type == "硬體開發":
        output += """- 設備費需說明必要性，優先考慮租用
- 打樣費用納入「消耗性器材」
- 認證測試列入「委託研究費」
"""
    elif project_type == "軟體開發":
        output += """- 雲端服務費需提供估算依據
- 軟體授權費可納入「消耗性器材」
- 人事費比例較高是正常的
"""
    elif project_type == "服務創新":
        output += """- 場地費需與服務內容相關
- 市場調查可列入「委託研究費」
- 可編列少量行銷推廣費用
"""
    else:
        output += """- 各項費用需附採購規劃說明
- 委外項目需說明必要性
- 差旅費需列明目的地和目的
"""

    output += """
---

> ⚠️ 此為建議分配，實際編列請依計畫需求調整
> 📖 詳細說明請參考：經費編列指南
"""

    return [TextContent(type="text", text=output)]

# ============================================
# Word 匯出功能
# ============================================


async def export_proposal_word(
    content: str,
    filename: str = "SBIR_計畫書",
    company_name: str = "",
    project_name: str = ""
) -> list[TextContent]:
    """
    將計畫書 Markdown 內容匯出為 Word 檔案
    """
    try:
        # 建立 Word 文件
        doc = Document()

        # 設定預設字型
        style = doc.styles['Normal']
        style.font.name = '標楷體'
        style.font.size = Pt(12)

        # 標題頁
        if project_name:
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(project_name)
            run.font.size = Pt(24)
            run.font.bold = True
            doc.add_paragraph()  # 空行

        if company_name:
            company = doc.add_paragraph()
            company.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = company.add_run(company_name)
            run.font.size = Pt(16)
            doc.add_paragraph()  # 空行

        # 分頁
        if project_name or company_name:
            doc.add_page_break()

        # 解析 Markdown 並轉換為 Word
        lines = content.split('\n')
        in_code_block = False

        for line in lines:
            # 處理程式碼區塊
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.left_indent = Inches(0.5)
                continue

            # 處理標題
            if line.startswith('# '):
                p = doc.add_paragraph(line[2:])
                run = p.runs[0]
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 128)
            elif line.startswith('## '):
                p = doc.add_paragraph(line[3:])
                run = p.runs[0]
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 128)
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:])
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True

            # 處理列表
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                p = doc.add_paragraph(text, style='List Number')

            # 處理分隔線
            elif line.strip() == '---':
                doc.add_paragraph('_' * 50)

            # 處理空行
            elif line.strip() == '':
                doc.add_paragraph()

            # 一般段落
            else:
                # 處理粗體 **text**
                if '**' in line:
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    doc.add_paragraph(line)

        # 儲存檔案 — 跨平台目錄偵測（~/Documents 在部分 Windows/Linux 不存在）
        home = Path.home()
        candidates = [home / "Desktop", home / "Documents", home / "文件", home / "桌面", home]
        output_dir = next((str(p) for p in candidates if p.exists()), str(home))
        output_path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(output_path)

        return [TextContent(
            type="text",
            text=f"""✅ **Word 檔案已成功匯出！**

📄 檔案位置：`{output_path}`

您可以：
1. 開啟檔案檢視內容
2. 根據需求調整格式
3. 直接用於 SBIR 申請送件

> 💡 提示：建議開啟後檢查格式，並補充圖表等視覺元素
"""
        )]

    except Exception as e:
        return [TextContent(
            type="text",
            text=f"❌ **匯出失敗**\n\n錯誤訊息：{str(e)}\n\n請確認：\n1. 已安裝 python-docx\n2. 有寫入權限到 ~/Documents"
        )]

# ============================================
# 計畫書完整度檢核功能
# ============================================


async def check_proposal(proposal_content: str, phase: str = "phase1") -> list[TextContent]:
    """
    檢核 SBIR 計畫書完整度
    這是「自我檢查工具」，不是「評審結果預測」
    """

    # 定義 Phase 1 檢核項目
    phase1_checks = [
        {
            "category": "基本資訊",
            "items": [
                {"name": "公司名稱", "keywords": ["公司", "股份有限", "有限公司"]},
                {"name": "計畫名稱", "keywords": ["計畫名稱", "計畫題目"]},
                {"name": "計畫期程", "keywords": ["期程", "月", "年"]},
            ]
        },
        {
            "category": "問題陳述",
            "items": [
                {"name": "產業痛點描述", "keywords": ["痛點", "問題", "挑戰", "困難", "需求"]},
                {"name": "現況說明", "keywords": ["現況", "目前", "現有", "傳統"]},
                {"name": "問題量化數據", "keywords": ["億", "萬", "%", "比例", "統計"]},
            ]
        },
        {
            "category": "創新內容",
            "items": [
                {"name": "創新點描述", "keywords": ["創新", "突破", "獨創", "首創", "原創"]},
                {"name": "與現有技術差異", "keywords": ["差異", "不同", "優於", "相較", "比較"]},
                {"name": "技術優勢說明", "keywords": ["優勢", "優點", "特色", "領先"]},
            ]
        },
        {
            "category": "市場分析",
            "items": [
                {"name": "目標市場描述", "keywords": ["目標市場", "客戶", "TA", "使用者"]},
                {"name": "市場規模（TAM/SAM/SOM）", "keywords": ["TAM", "SAM", "SOM", "市場規模", "產值"]},
                {"name": "商業模式", "keywords": ["商業模式", "獲利", "營收", "收費"]},
            ]
        },
        {
            "category": "技術可行性",
            "items": [
                {"name": "技術方案說明", "keywords": ["技術", "方法", "架構", "系統"]},
                {"name": "前期驗證成果", "keywords": ["驗證", "測試", "實驗", "前期", "雛型"]},
                {"name": "風險評估", "keywords": ["風險", "挑戰", "困難"]},
            ]
        },
        {
            "category": "團隊介紹",
            "items": [
                {"name": "團隊成員", "keywords": ["團隊", "成員", "人員"]},
                {"name": "相關經驗", "keywords": ["經驗", "經歷", "背景", "專長"]},
                {"name": "分工規劃", "keywords": ["分工", "負責", "職責"]},
            ]
        },
        {
            "category": "執行計畫",
            "items": [
                {"name": "工作項目", "keywords": ["工作", "項目", "任務"]},
                {"name": "時程規劃", "keywords": ["時程", "進度", "甘特", "月"]},
                {"name": "查核點", "keywords": ["查核", "里程碑", "KPI", "指標"]},
            ]
        },
        {
            "category": "經費規劃",
            "items": [
                {"name": "人事費", "keywords": ["人事費", "薪資", "人力"]},
                {"name": "材料費/設備費", "keywords": ["材料", "設備", "器材", "耗材"]},
                {"name": "其他費用", "keywords": ["委託", "差旅", "管理費"]},
            ]
        },
    ]

    # 執行檢核
    content_lower = proposal_content.lower()
    results = []
    total_items = 0
    passed_items = 0

    for category in phase1_checks:
        category_results = {
            "name": category["category"],
            "items": []
        }

        for item in category["items"]:  # type: ignore
            total_items += 1
            # 檢查是否包含關鍵字（不區分大小寫）
            found = any(keyword.lower() in content_lower for keyword in item["keywords"])  # type: ignore
            if found:
                passed_items += 1
                status = "✅"
            else:
                status = "❌"

            category_results["items"].append({  # type: ignore
                "name": item["name"],  # type: ignore
                "status": status,
                "found": found
            })

        results.append(category_results)

    # 格式化輸出
    output = f"""# 📋 SBIR 計畫書完整度檢核

> ⚠️ **重要提醒**：這是「自我檢查工具」，用來確認計畫書是否涵蓋必要內容。
> 檢核結果 **不代表審查結果預測**，最終通過與否取決於審查委員評估。

---

## 檢核結果摘要

**完整度**：{passed_items}/{total_items} 項目已涵蓋（{int(passed_items/total_items*100)}%）

"""

    for category in results:
        category_passed = sum(1 for item in category["items"] if item["found"])  # type: ignore
        category_total = len(category["items"])  # type: ignore

        if category_passed == category_total:
            category_status = "✅"
        elif category_passed == 0:
            category_status = "❌"
        else:
            category_status = "⚠️"

        output += f"### {category_status} {category['name']} ({category_passed}/{category_total})\n\n"

        for item in category["items"]:  # type: ignore
            output += f"- {item['status']} {item['name']}\n"  # type: ignore

        output += "\n"

    # 添加建議
    missing_items = [
        f"- {item['name']}"  # type: ignore
        for category in results
        for item in category["items"]  # type: ignore
        if not item["found"]  # type: ignore
    ]

    if missing_items:
        output += """---

## 💡 建議補強項目

以下項目可能需要補充或加強：

"""
        for item in missing_items[:10]:  # 最多顯示 10 項
            output += f"{item}\n"

        if len(missing_items) > 10:
            output += f"\n（還有 {len(missing_items) - 10} 項未列出）\n"
    else:
        output += """---

## 🎉 恭喜！

您的計畫書涵蓋了所有必要項目。建議進一步優化：
- 確認各項內容的深度和具體性
- 補充量化數據和佐證資料
- 請他人審閱並給予回饋
"""

    output += """
---

📖 需要更多指引？請說「搜尋 [關鍵字]」查詢知識庫
"""

    return [TextContent(type="text", text=output)]

# ============================================
# ROI 計算工具
# ============================================


async def call_calculate_roi(
    subsidy_amount: float,
    industry: str = "製造業",
    phase: str = "phase1",
    company_revenue: float = 0
) -> list[TextContent]:
    """計算建議的產值目標"""
    try:
        from roi_calculator import calculate_roi, format_roi_report

        result = calculate_roi(
            subsidy_amount=subsidy_amount,
            phase=phase,
            industry=industry,
            project_duration=12 if phase == "phase1" else 24,
            company_revenue=company_revenue
        )

        report = format_roi_report(result)

        return [TextContent(type="text", text=report)]

    except Exception as e:
        return [TextContent(
            type="text",
            text=f"❌ ROI 計算失敗：{str(e)}"
        )]


async def call_validate_roi(
    subsidy_amount: float,
    expected_revenue_3years: float,
    industry: str = "製造業",
    phase: str = "phase1"
) -> list[TextContent]:
    """驗證產值是否合理"""
    try:
        from roi_calculator import validate_roi, format_validation_report

        result = validate_roi(
            subsidy_amount=subsidy_amount,
            expected_revenue_3years=expected_revenue_3years,
            industry=industry,
            phase=phase
        )

        report = format_validation_report(result)

        return [TextContent(type="text", text=report)]

    except Exception as e:
        return [TextContent(
            type="text",
            text=f"❌ ROI 驗證失敗：{str(e)}"
        )]

# ============================================
# 主程式入口
# ============================================


async def main():
    """啟動 MCP Server"""
    from mcp.server.stdio import stdio_server

    async with stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            app.create_initialization_options()
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
